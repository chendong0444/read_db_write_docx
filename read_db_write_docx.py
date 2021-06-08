from docx import table
import yaml
import json
import os
from docx import Document


BASE_PATH = os.path.dirname(__file__)
with open(os.path.join(BASE_PATH, 'config.yml'),encoding='utf8') as f:
    conf = yaml.load(f, Loader=yaml.FullLoader)


def get_tbl_struct(tbl_name): # 获取表结构，返回
    db_conf = conf['db_info']
    db_type = db_conf['db_type']
    if db_type == 'mysql':
        return get_mysql_tbl_struct(tbl_name)
    elif db_type == 'oracle':
        return get_oracle_tbl_struct(tbl_name)
    elif db_type == 'sqlserver':
        return get_sqlserver_tbl_struct(tbl_name)
    else:
        raise Exception('不支持的数据库类型',db_type)


def get_sqlserver_tbl_struct(tbl_name):
    """通过information_schema.COLUMNS, 读取表结构信息
    """
    # 链接数据库
    import pymssql
    db_conf = conf['db_info']
    db_name = db_conf['db']
    conn = pymssql.connect(host=db_conf['host'], port=int(db_conf['port']), user=db_conf['user'],
                           password=db_conf['password'], database=db_conf['db'], charset=db_conf['charset'])
    # 第二步：创建游标对象
    cursor = conn.cursor()  # cursor当前的程序到数据之间连接管道
    # 第三步：组装sql语句
    #sql = f"select column_name,column_type,data_type,CHARACTER_MAXIMUM_LENGTH,is_nullable,column_comment from `information_schema`.`COLUMNS`  where `table_name` = '{tbl_name}' and `table_schema` = '{db_name}' order by ordinal_position"

    sql=f"""
                    SELECT
                字段名    =c.name  ,
                类型      =y.Name ,
                字节      =c.max_length ,
                主键      =case when exists(SELECT 1 FROM sysobjects where xtype='PK' and parent_obj=t.object_id and name in (
                          SELECT name FROM sysindexes WHERE indid in(
                          SELECT indid FROM sysindexkeys WHERE id =t.object_id AND colid=c.column_id))) then '√' else '' end,
                是否可为空 =case when c.is_nullable='1' then 'Yes' else 'No'  end,
                字段说明   =SUBSTRING(cast(ep.[value] as nvarchar(200)),1,charindex('|',cast(ep.[value] as nvarchar(200)))-1)
        FROM sys.tables AS t INNER JOIN sys.columns  AS c ON t.object_id = c.object_id
        LEFT JOIN  sys.extended_properties AS ep ON ep.major_id = c.object_id AND ep.minor_id = c.column_id
        LEFT JOIN  sys.Types AS y ON y.User_Type_ID=c.User_Type_ID
        WHERE ep.class =1   AND t.name='"""+tbl_name+"""'

        """
    # 第四步：执行sql语句
    cursor.execute(sql)
    # 从游标中取出所有记录放到一个序列中并关闭游标
    result = cursor.fetchall()
    # 查询表的结构
    fields = list(result)
    # 列名	数据类型	字段类型	长度	是否为空	默认值	备注
    # print('|列名|数据类型|字段类型|长度|是否为空|备注|'.replace('|',','))
    data = ['字段名|类型|字节|主键|是否可为空|字段说明'.split('|')]
    # print('|--|--|--|--|--|--|')
    for f in fields:
        s = [str(i) if i is not None else '' for i in f]
        data.append(s)
    # 关闭游标
    cursor.close()
    conn.close()
    return data


def get_tbl_names():
    db_conf = conf['db_info']
    db_type = db_conf['db_type']
    if db_type == 'mysql':
        return get_mysql_tbl_names()
    elif db_type == 'oracle':
        return get_oracle_tbl_names()
    elif db_type == 'sqlserver':
        raise Exception('sqlserver还不支持自动获取表名，需要自己手工定义word_def.tables')
    else:
        raise Exception('不支持的数据库类型', db_type)


def get_oracle_tbl_names():
    tbl_names = []
    # 链接数据库
    import cx_Oracle
    db_conf = conf['db_info']

    conn = cx_Oracle.connect(
        '%s/%s@%s:%s/%s' % (db_conf['user'], db_conf['password'], db_conf['host'], db_conf['port'], db_conf['db']))

    # 第二步：创建游标对象
    cursor = conn.cursor()  # cursor当前的程序到数据之间连接管道
    # 第三步：组装sql语句
    sql = 'select a.table_name, b.comments ' \
          'from user_tables a inner join user_tab_comments b ' \
          'on a.table_name=b.table_name ' \
          'order by a.table_name'
    # 第四步：执行sql语句
    cursor.execute(sql)
    # 从游标中取出所有记录放到一个序列中并关闭游标
    result = cursor.fetchall()

    for item in list(result):
        if item[1]:
            tbl_names.append('__'.join(item))
        else:
            tbl_names.append(item[0])

    # 关闭游标
    cursor.close()
    conn.close()
    return tbl_names


def get_mysql_tbl_names():
    """通过information_schema.TABLES, 读取表名，格式为{表名}__{备注}
    """
    # 链接数据库
    import pymysql
    db_conf = conf['db_info']
    db_name = db_conf['db']
    print("db_conf=",db_conf)
    conn = pymysql.connect(host=db_conf['host'], port=int(db_conf['port']), user=db_conf['user'],
                           password=db_conf['password'], db=db_conf['db'], charset=db_conf['charset'])
                           
    print("end db_conf=",conn)
    # 第二步：创建游标对象
    cursor = conn.cursor()  # cursor当前的程序到数据之间连接管道
    # 第三步：组装sql语句
    sql = f"select TABLE_NAME,TABLE_COMMENT from `information_schema`.`TABLES` as t where t.TABLE_SCHEMA='{db_conf['db']}' order by TABLE_NAME"
    # sql = f"select column_name,column_type,data_type,CHARACTER_MAXIMUM_LENGTH,is_nullable,column_comment from `information_schema`.`COLUMNS`  where `table_name` = '{tbl_name}' and `table_schema` = '{db_name}' order by ordinal_position"
    # 第四步：执行sql语句
    cursor.execute(sql)
    # 从游标中取出所有记录放到一个序列中并关闭游标
    result = cursor.fetchall()
    # 查询表的结构
    rows = list(result)
    data = []
    # print('|--|--|--|--|--|--|')
    for row in rows:
        data.append(f'{row[0]}__{row[1]}')
    # 关闭游标
    cursor.close()
    conn.close()
    print(data)
    return data


def get_mysql_tbl_struct(tbl_name):
    """通过information_schema.COLUMNS, 读取表结构信息
    """
    # 链接数据库
    import pymysql
    db_conf = conf['db_info']
    db_name = db_conf['db']
    print("db_conf=",db_conf)
    conn = pymysql.connect(host=db_conf['host'], port=int(db_conf['port']), user=db_conf['user'],
                           password=db_conf['password'], db=db_conf['db'], charset=db_conf['charset'])

    print("end db_conf=",conn)
    # 第二步：创建游标对象
    cursor = conn.cursor()  # cursor当前的程序到数据之间连接管道
    # 第三步：组装sql语句
    sql = f"select column_name,column_type,data_type,CHARACTER_MAXIMUM_LENGTH,is_nullable,column_comment from `information_schema`.`COLUMNS`  where `table_name` = '{tbl_name}' and `table_schema` = '{db_name}' order by ordinal_position"
    # 第四步：执行sql语句
    cursor.execute(sql)
    # 从游标中取出所有记录放到一个序列中并关闭游标
    result = cursor.fetchall()
    # 查询表的结构
    fields = list(result)
    # 列名	数据类型	字段类型	长度	是否为空	默认值	备注
    # print('|列名|数据类型|字段类型|长度|是否为空|备注|'.replace('|',','))
    data = ['列名|数据类型|字段类型|长度|是否为空|备注'.split('|')]
    # print('|--|--|--|--|--|--|')
    for f in fields:
        s = [str(i) if i is not None else '' for i in f]
        data.append(s)
    # 关闭游标
    cursor.close()
    conn.close()
    return data


def get_oracle_tbl_struct(tbl_name):
    # 链接数据库
    import cx_Oracle
    db_conf = conf['db_info']

    conn = cx_Oracle.connect('%s/%s@%s:%s/%s' % (db_conf['user'], db_conf['password'], db_conf['host'], db_conf['port'], db_conf['db']))

    # 第二步：创建游标对象
    cursor = conn.cursor()  # cursor当前的程序到数据之间连接管道
    # 第三步：组装sql语句
    sql = f"""
            select "字段名称","数据类型","长度","是否为空","字段说明","默认值",Concat(CONCAT("主键表名",'.'),"主键列名") as "外键约束" from (
                SELECT t1.Table_Name AS "表名称",
                t3.comments AS "表说明",
                t1.Column_Name AS "字段名称",
                t1.Data_Type AS "数据类型",
                t1.Data_Length AS "长度",
                t1.NullAble AS "是否为空",
                t2.Comments AS "字段说明",
                t1.Data_Default "默认值",
                t4.主键表名 AS 主键表名,
                t4.主键列名 as 主键列名
                FROM cols t1 left join user_col_comments t2
                on t1.Table_name=t2.Table_name and t1.Column_Name=t2.Column_Name
                left join user_tab_comments t3
                on t1.Table_name=t3.Table_name
                left join (
                    select distinct aa.table_name 外键表名,aa.column_name 外键列名,bb.table_name 主键表名,bb.column_name 主键列名
                    from
                    (select a.constraint_name,b.table_name,b.column_name,a.r_constraint_name
                    from user_constraints a, user_cons_columns b
                    WHERE a.constraint_type='R'
                    and a.constraint_name=b.constraint_name
                    ) aa,
                    (select a.r_constraint_name,b.table_name,b.column_name
                    from user_constraints a, user_cons_columns b where a.constraint_type='R' and
                    a.r_constraint_name=b.constraint_name)
                    bb
                    where aa.r_constraint_name=bb.r_constraint_name
                ) t4
                on t1.Table_name=t4.外键表名 and t1.Column_Name=t4.外键列名 and t3.Table_name = t4.外键表名 and t2.Table_name = t4.外键表名
                WHERE NOT EXISTS ( SELECT t4.Object_Name FROM User_objects t4
                    WHERE t4.Object_Type='TABLE'
                    AND t4.Temporary='Y'
                    AND t4.Object_Name=t1.Table_Name )
                ORDER BY t1.Table_Name, t1.Column_ID
            )k where k."表名称" = '{tbl_name}'
        """
    # 第四步：执行sql语句
    cursor.execute(sql)
    # 从游标中取出所有记录放到一个序列中并关闭游标
    result = cursor.fetchall()
    # 查询表的结构
    fields = list(result)
    data = ['字段名称|数据类型|长度|是否可为空|字段说明|默认值|外键约束'.split('|')]
    # print('|--|--|--|--|--|--|')
    for field in fields:
        arr = list(field)
        if arr[-1] == '.':
            arr[-1] = ''
        data.append(arr)
    print(data)
    # 关闭游标
    cursor.close()
    conn.close()
    return data


def insert_after_paragraph(_p1, _p2):
    """在docx中做插入操作
    """
    p1 = _p1._tbl if isinstance(_p1, table.Table) else _p1._p
    p2 = _p2._tbl if isinstance(_p2, table.Table) else _p2._p
    p1.addnext(p2)


def find_anchor_paragraph(anchor_text):
    # 找到定义锚点的段落
    paragraphs = document.paragraphs
    for paragraph in paragraphs:
        if anchor_text == paragraph.text:
            return paragraph
    raise Exception('没有找到锚点:'+anchor_text)


def get_next_level_style(paragraph, step=1):
    # 自动查找下一级标题的style
    new_style = paragraph.style.name[:-1] + \
        str(int(paragraph.style.name[-1])+step)
    return new_style


def createDocxTable(items, document):
    """创建一个docx格式的表格,column_len_def表示表格的列宽定义
    """

    # 表字段长度
    column_len_def = [2.8, 2.5, 1.5, 1.5, 2.5, 2.0, 2.0]

    # add table ------------------
    colunm_len = len(items[0])
    table = document.add_table(len(items), len(items[0]))
    from docx.shared import Cm

    def set_column_width(column, width):
        column.width = width
        for cell in column.cells:
            cell.width = width

    if column_len_def:
        table.autofit = False
        table.allow_autofit = False
        for i in range(len(column_len_def)):
            width = Cm(column_len_def[i])
            set_column_width(table.columns[i], width)

    heading_cells = table.rows[0].cells
    for i in range(colunm_len):
        heading_cells[i].text = str(items[0][i])

    # add a data row for each item
    for j in range(1, len(items)):
        cells = table.rows[j].cells
        for i in range(colunm_len):
            cells[i].text = str(items[j][i])

    table.style = 'Table Grid'
    return table


def read_db_write_docx():
    # 往word中写入表格内容,可以支持多个段落定义
    for section in conf['word_def']:
        anchor = section['anchor']
        tables = section.get('tables',[])
        if not tables:
            tables = get_tbl_names()
        p = find_anchor_paragraph(anchor)
        new_style = get_next_level_style(p)
        for t in tables:
            # 创建并插入标题
            if '__' in t:
                tbl_name, tbl_name_chs = t.split('__')  # 表名和中文名要用下划线分开
            else:
                tbl_name, tbl_name_chs = t, ''
            if tbl_name_chs:
                display_name = f'{tbl_name}({tbl_name_chs})'
            else:
                display_name = tbl_name
            x = document.add_paragraph(display_name, style=new_style)
            insert_after_paragraph(p, x)

            tbl_struct_info = get_tbl_struct(tbl_name)  # 获取表结构，返回

            # 创建并插入表结构
            docx_t = createDocxTable(tbl_struct_info, document)
            insert_after_paragraph(x, docx_t)
            p = docx_t


if __name__ == '__main__':
    document = Document(conf['template'])
    read_db_write_docx()
    document.save(conf['output'])

# https://github.com/python-openxml/python-docx/issues/156 在一段后面插入表格
# 引用，https://github.com/python-openxml/python-docx/issues/823  如何根据文本内容找到某一段
# https://github.com/python-openxml/python-docx/issues/33  在docx中删除一个段落